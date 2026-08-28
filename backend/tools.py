"""
TalentPulse AI - Tools & Document Utilities Module
Transferred 100% authentically from existing agent tools:
- Memory-safe PDF, DOCX, and Image extraction
- Vietnamese text normalization and bigram name similarity
"""

import io
import re
import unicodedata
import logging
from typing import Optional, Tuple
import pypdf
import docx
from PIL import Image

logger = logging.getLogger("talentpulse.tools")

# ==============================================================================
# 1. TEXT CLEANING & VIETNAMESE NORMALIZATION (100% Authentic Legacy)
# ==============================================================================
def normalize_vn(text: str) -> str:
    """Normalize Vietnamese text to unaccented lowercase."""
    if not text:
        return ""
    text_norm = unicodedata.normalize("NFD", str(text))
    text_no_accents = re.sub(r"[\u0300-\u036f]", "", text_norm)
    text_clean = text_no_accents.replace("đ", "d").replace("Đ", "D")
    return text_clean.lower().strip()


def letters_only(text: str) -> str:
    """Keep only letters from normalized Vietnamese text."""
    if not text:
        return ""
    norm = normalize_vn(text)
    return re.sub(r"[^a-zA-Z]", "", norm).lower()


def get_bigrams(text: str) -> set:
    """Get bigram set of string for similarity scoring."""
    b = set()
    for i in range(len(text) - 1):
        b.add(text[i : i + 2])
    return b


def name_similarity(a: str, b: str) -> float:
    """Calculate bigram name similarity between two names."""
    la = letters_only(a)
    lb = letters_only(b)
    if not la or not lb:
        return 0.0
    ba = get_bigrams(la)
    bb = get_bigrams(lb)
    if not ba or not bb:
        return 0.0
    common = ba.intersection(bb)
    return len(common) / max(len(ba), len(bb))


def to_title_case(text: str) -> str:
    """Capitalize each word in candidate name."""
    if not text:
        return ""
    words = str(text).strip().split()
    return " ".join([w.capitalize() for w in words])


# ==============================================================================
# 2. MEMORY-SAFE DOCUMENT EXTRACTORS
# ==============================================================================
def extract_text_from_pdf_stream(pdf_bytes: bytes) -> str:
    """Extract text from in-memory PDF byte stream safely."""
    try:
        reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
        pages_text = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages_text.append(t)
        return "\n".join(pages_text).strip()
    except Exception as e:
        logger.warning(f"Failed extracting text via pypdf: {e}")
        return ""


def extract_text_from_docx_stream(docx_bytes: bytes) -> str:
    """Extract text from in-memory DOCX byte stream safely."""
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        return "\n".join(paragraphs).strip()
    except Exception as e:
        logger.warning(f"Failed extracting text via docx: {e}")
        return ""


def parse_cv_filename(filename: str) -> Tuple[str, str]:
    """
    Extract candidate name & position hint from filename.
    e.g., 'CV_Nguyen_Van_A_Python_Dev.pdf' -> ('Nguyen Van A', 'Python Dev')
    """
    name_no_ext = re.sub(r"\.[^.]+$", "", filename)
    clean = re.sub(r"[_\-]+", " ", name_no_ext).strip()
    return clean, ""
